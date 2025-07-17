#pip install python-docx
# pip install --upgrade openai
import re
from datetime import datetime
#  From the python-docx library, used to read and write Word .docx files.
from docx import Document

# Imports the OpenAI v1.x SDK client for connecting to GPT models
from openai import OpenAI

# --- Configuration ---
TEMPLATE_PATH = "C:/Users/ishit/OneDrive/Documents/Python/writeup_generator/Writeup_template.docx"
PROMPT_PATH = "C:/Users/ishit/OneDrive/Documents/Python/writeup_generator/issue_prompt.txt"
API_KEY = ""  # Replace with your actual API key

client = OpenAI(api_key=API_KEY)

# --- Load issue prompt ---
with open(PROMPT_PATH, "r", encoding="utf-8") as f:
    raw_prompt = f.read()

# --- Extract PIM number ---
match = re.search(r'PIM[-\s#]*([0-9]+)', raw_prompt, re.IGNORECASE)
pim_number = match.group(1) if match else datetime.now().strftime("%Y%m%d_%H%M")

# --- Load Word Template ---
doc = Document(TEMPLATE_PATH)

# --- Find All Placeholders in Template ---
# Scans through all paragraphs in the document and extracts anything in <<...>> as a placeholder 
# field name.
placeholders = set()
for para in doc.paragraphs:
    matches = re.findall(r"<<(.*?)>>", para.text)
    placeholders.update(matches)

# --- Ask ChatGPT to Generate Text for Each Placeholder ---
def get_field_content(field, full_context):
    system_msg = (
        "You are generating structured technical RCA content to fill fields in an O&M document."
        " Each field must be written professionally and concisely."
    )
    user_msg = f"""Based on the following issue details, generate content for the field: "{field}"

Issue Details:
{full_context}

Only return clean text for the field "{field}".
"""
 
# This is the ChatGPT API call, using the v1 SDK.
# It sends a message sequence (system + user) and requests a structured, formal reply.
# temperature=0.5: lower value means more consistent, formal outputs.

    completion = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg}
        ],
        temperature=0.5,
    )
    return completion.choices[0].message.content.strip()

# --- Replace Placeholders in Document ---
# for para in doc.paragraphs:
#     for ph in placeholders:
#         placeholder_tag = f"<<{ph}>>"
#         if placeholder_tag in para.text:
#             print(f"Generating content for: {ph}")
#             ai_response = get_field_content(ph, raw_prompt)
#             para.text = para.text.replace(placeholder_tag, ai_response)


raw_prompt_lower = raw_prompt.lower()
for para in doc.paragraphs:
    for ph in placeholders:
        placeholder_tag = f"<<{ph}>>"

        if placeholder_tag not in para.text:
            continue
        # Only process if placeholder keyword is mentioned in the prompt
        if ph.replace("_", " ").lower() not in raw_prompt_lower:
            print(f"Skipping '{ph}' — not found in prompt")
            continue

        print(f"Generating content for: {ph}")
        ai_response = get_field_content(ph, raw_prompt)
        para.text = para.text.replace(placeholder_tag, ai_response)


# --- Save Final Write-up File ---
output_file = f"Writeup_PIM-{pim_number}.docx"
doc.save(output_file)
print(f" Generated write-up saved to: {output_file}")
