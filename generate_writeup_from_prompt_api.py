import re
import requests
import json
from datetime import datetime
from docx import Document

# --- Configuration ---
TEMPLATE_PATH = "C:/Users/ishit/OneDrive/Documents/Python/writeup_generator/Writeup_template.docx"
PROMPT_PATH = "C:/Users/ishit/OneDrive/Documents/Python/writeup_generator/issue_prompt.txt"
API_KEY = "Key.."  # Replace with your actual API key
API_URL = "https://api.openai.com/v1/chat/completions"

# --- Load issue prompt ---
with open(PROMPT_PATH, "r", encoding="utf-8") as f:
    raw_prompt = f.read()

# --- Extract PIM number ---
match = re.search(r'PIM[-\s#]*([0-9]+)', raw_prompt, re.IGNORECASE)
pim_number = match.group(1) if match else datetime.now().strftime("%Y%m%d_%H%M")

# --- Load Word Template ---
doc = Document(TEMPLATE_PATH)

# --- Find All Placeholders in Template ---
placeholders = set()
for para in doc.paragraphs:
    matches = re.findall(r"<<(.*?)>>", para.text)
    placeholders.update(matches)

# --- Function to Call ChatGPT API ---
def get_field_content(field, full_context):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    system_msg = (
        "You are generating structured technical RCA content to fill fields in an O&M document. "
        "Each field must be written professionally and concisely."
    )
    user_msg = f"""Based on the following issue details, generate content for the field: "{field}"

Issue Details:
{full_context}

Only return clean text for the field "{field}".
"""

    data = {
        "model": "gpt-4",
        "messages": [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg}
        ],
        "temperature": 0.5
    }

    response = requests.post(API_URL, headers=headers, data=json.dumps(data))

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"].strip()
    else:
        print(f"Error from API: {response.status_code} - {response.text}")
        return f"[Error generating {field}]"

# --- Replace Placeholders If Mentioned in Prompt ---
raw_prompt_lower = raw_prompt.lower()

for para in doc.paragraphs:
    for ph in placeholders:
        placeholder_tag = f"<<{ph}>>"

        if placeholder_tag not in para.text:
            continue

        if ph.replace("_", " ").lower() not in raw_prompt_lower:
            print(f"⏭️ Skipping '{ph}' — not found in prompt")
            continue

        print(f" Generating content for: {ph}")
        ai_response = get_field_content(ph, raw_prompt)
        para.text = para.text.replace(placeholder_tag, ai_response)

# --- Save Final Document ---
output_file = f"C:/Users/ishit/OneDrive/Documents/Python/writeup_generator/Writeup_PIM-{pim_number}.docx"
doc.save(output_file)
print(f" Generated write-up saved to: {output_file}")
